from django.http import HttpResponse
from django_xhtml2pdf.utils import generate_pdf
from docx import Document
from lxml import etree
from openpyxl import Workbook

# from openpyxl.utils import get_column_letter
from openpyxl.styles import Font
from openpyxl.styles import PatternFill


def render_xls_igce(context, **kwargs):
    """
    """

    workbook = Workbook()
    filename = kwargs.get("filename")
    item = context["item"]
    bold = Font(bold=True)

    sheet1 = workbook.active
    sheet1.title = "Instructions"
    # https://stackoverflow.com/a/14450572
    sheet1.append(["Instructions".upper()])
    sheet1.column_dimensions["B"].width = 111.7
    sheet1["A1"].font = bold
    sheet1.merge_cells("A1:Z1")
    sheet1.append(
        [
            "",
            "This template is being provided as a tool to assist the acquisition workforce in developing an IGCE for a Firm Fixed Price Product or Service.",
        ]
    )
    sheet1.append([])
    sheet1.append(
        [
            "1.",
            "The exact amount of a vendor’s quote must NOT be used as the basis for a program office’s IGCE.",
        ]
    )
    sheet1.append(
        [
            "2.",
            "The program office must conduct all necessary research to compile an accurate and complete IGCE, independent of the vendor’s pricing/cost information.",
        ]
    )
    sheet1.append(
        [
            "3.",
            "The program office should use the results of the market reaearch to substaniate the IGCE.",
        ]
    )
    sheet1.append(
        [
            "4.",
            "The program office provide a narrative to document the basis of the IGCE.",
        ]
    )
    sheet2 = workbook.create_sheet(title="FFP IGCE")
    sheet2.append(["FFP IGCE TEMPLATE"])
    sheet2.append(["TITLE:"])
    sheet2.append(["Detailed Price Summary"])
    sheet2.append(
        [
            "Contract Line Item Description",
            "Estimate 1",
            "",
            "",
            "",
            "Estimate 2",
            "",
            "",
            "",
            "Estimate 3",
        ]
    )
    sheet2.append(
        [
            "",
            "Quantity",
            "Unit",
            "Unit Price",
            "Total Price",
            "Quantity",
            "Unit",
            "Unit Price",
            "Total Price",
            "Quantity",
            "Unit",
            "Unit Price",
            "Total Price",
        ]
    )
    sheet2["A1"].font = bold
    sheet2["A2"].font = bold
    sheet2["A3"].font = bold
    sheet2["A4"].font = bold

    sheet2["B5"].font = bold
    sheet2["C5"].font = bold
    sheet2["D5"].font = bold
    sheet2["E5"].font = bold

    sheet2["F5"].font = bold
    sheet2["G5"].font = bold
    sheet2["H5"].font = bold
    sheet2["I5"].font = bold

    sheet2["J5"].font = bold
    sheet2["K5"].font = bold
    sheet2["L5"].font = bold
    sheet2["M5"].font = bold

    sheet2.column_dimensions["A"].width = 48

    sheet2.merge_cells("B4:E4")
    sheet2.merge_cells("F4:I4")
    sheet2.merge_cells("J4:M4")

    # https://stackoverflow.com/a/50209914
    sheet2["B4"].fill = PatternFill(
        start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
    )
    sheet2["F4"].fill = PatternFill(
        start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
    )
    sheet2["J4"].fill = PatternFill(
        start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
    )

    sheet2["B4"].font = bold
    sheet2["F4"].font = bold
    sheet2["J4"].font = bold

    entries = []
    for entry in item.time_set.all():
        entries.append(entry.task.name)
        entries.append(entry.description)
        entries.append(entry.hours)
        entries.append(entry.task.rate)

    entries.insert(0, "")
    sheet2.append(entries)

    sheet2["E" + str(sheet2.max_row)].fill = PatternFill(
        start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
    )
    sheet2["I" + str(sheet2.max_row)].fill = PatternFill(
        start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
    )
    sheet2["M" + str(sheet2.max_row)].fill = PatternFill(
        start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
    )

    sheet2.append(["Line Item Subtotal"])

    sheet2["B" + str(sheet2.max_row)].fill = PatternFill(
        start_color="00008B", end_color="00008B", fill_type="solid"
    )
    sheet2["C" + str(sheet2.max_row)].fill = PatternFill(
        start_color="00008B", end_color="00008B", fill_type="solid"
    )
    sheet2["D" + str(sheet2.max_row)].fill = PatternFill(
        start_color="00008B", end_color="00008B", fill_type="solid"
    )
    sheet2["E" + str(sheet2.max_row)].fill = PatternFill(
        start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
    )

    response = HttpResponse(content_type="xlsx")
    response["Content-Disposition"] = "attachment; filename=%s" % filename
    workbook.save(response)
    return response


def render_doc(context, **kwargs):
    """
    """

    # https://python-docx.readthedocs.io/en/latest/#what-it-can-do
    document = Document()
    filename = kwargs.get("filename")
    item = context["item"]
    document.add_heading(item.title, 0)
    root = etree.fromstring(item.text)
    for elem in root.iter():
        if elem.tag == "h1":
            document.add_heading(elem.text, level=1)
            document.add_paragraph("")
        if elem.tag == "h2":
            document.add_heading(elem.text, level=2)
            document.add_paragraph("")
        if elem.tag == "h3":
            document.add_heading(elem.text, level=3)
            document.add_paragraph("")
        if elem.tag == "h4":
            document.add_heading(elem.text, level=4)
            document.add_paragraph("")
        if elem.tag == "h5":
            document.add_heading(elem.text, level=5)
            document.add_paragraph("")
        if elem.tag == "h6":
            document.add_heading(elem.text, level=6)
            document.add_paragraph("")
        if elem.tag == "p":
            document.add_paragraph(elem.text)
            document.add_paragraph("")

    response = HttpResponse(content_type="docx")
    response["Content-Disposition"] = "attachment; filename=%s" % filename
    document.save(response)
    return response


def render_pdf(context, **kwargs):
    """
    """

    filename = kwargs.get("filename")
    template = kwargs.get("template")
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = "filename=%s" % filename
    return generate_pdf(template, context=context, file_object=response)


def render_xls(context, **kwargs):
    """
    """

    # https://openpyxl.readthedocs.io/en/stable/usage.html#write-a-workbook
    workbook = Workbook()
    filename = kwargs.get("filename")
    item = context["item"]
    sheet1 = workbook.active
    sheet1.title = "Invoice"
    for entry in item.time_set.all():
        sheet1.append(
            [
                entry.date,
                entry.task.name,
                entry.description,
                entry.hours,
                entry.task.rate,
            ]
        )

    response = HttpResponse(content_type="xlsx")
    response["Content-Disposition"] = "attachment; filename=%s" % filename
    workbook.save(response)
    return response
